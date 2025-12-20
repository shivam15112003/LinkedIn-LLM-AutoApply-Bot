import sys
import os
from docx import Document

def txt_to_docx(txt_path: str, docx_path: str) -> None:
    doc = Document()
    with open(txt_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.rstrip("\n")
            if line.strip() == "":
                # blank line -> new paragraph
                doc.add_paragraph()
            else:
                doc.add_paragraph(line)
    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python txt_to_docx.py input.txt output.docx")
        sys.exit(1)

    inp = sys.argv[1]
    out = sys.argv[2]

    if not os.path.exists(inp):
        print(f"Input file not found: {inp}")
        sys.exit(1)

    txt_to_docx(inp, out)
